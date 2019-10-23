import os
from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

class WriteDocx:
    'word操作类，根据不同的输入内容，生成word文件'
    def __init__(self, name='xxx'):
        self.name = name
        self.doc = Document()

    def add_title(self, text):
        '书写标题'
        self.doc.add_heading(text)

    def add_subtitle(self, text):
        '书写副标题'
        self.doc.add_paragraph(text, style='Subtitle')

    def add_body(self, text):
        '书写正文内容'
        self.doc.add_paragraph(text)

    def save_file(self):
        self.doc.save(os.path.join(BASE_DIR, 'db', f'{self.name}.docx'))
        print(f'名为{self.name}.docx的文件已成功生成。')

def main(name='xxx'):
    writedocx = WriteDocx(name)
    title = input('请输入标题：')
    subtitle = input('请输入副标题：')
    body = input('请输入正文：')
    writedocx.add_title(title)
    writedocx.add_subtitle(subtitle)
    writedocx.add_body(body)
    writedocx.save_file()

    writedocx.add_title(title)

if __name__ == "__main__":
    main()


    